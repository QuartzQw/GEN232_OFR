from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.text.run import *
import json
from docx.enum.text import WD_LINE_SPACING

def create_choice_inform(runChoices, choiceDict):
    choiceType = choiceDict["choiceType"]
    choice = choiceDict["choice"] 
    # varies choice behavior for different input area in docx 
    if choiceType == "checkBox":
        runChoices.add_picture("../template/circle.jpg", width=Inches(0.12))
        runChoices.add_text(f" {choice}")
        runChoices.add_text("            ")
    elif choiceType == "checkBoxWithText":
        runChoices.add_picture("../template/circle.jpg", width=Inches(0.12))
        runChoices.add_text(f" {choice} ")
        runChoices.add_picture("../template/grayBox.jpg", width=Inches(1), height = Inches(0.25))
        runChoices.add_text("            ")
    elif choiceType == "text":
        runChoices.add_text(f" {choice} ")
        runChoices.add_picture("../template/grayBox.jpg", width=Inches(1), height = Inches(0.25))
        runChoices.add_text("    ")
    elif choiceType == "longText":
        runChoices.add_text(f" {choice} ")
        runChoices.add_picture("../template/grayBox.jpg", width=Inches(3), height = Inches(0.25))
        runChoices.add_text("    ")
    elif choiceType == "number":
        runChoices.add_text(f" {choice} ")
        runChoices.add_picture("../template/grayBox.jpg", width=Inches(1), height = Inches(0.25))
        runChoices.add_text(f' {choiceDict["unit"]}   ')
    

def create_question_in_form(question):

    # add question text to document
    questionLine = document.add_paragraph(f'{question["question"]}')
    questionLine.style.font.size = Pt(16)
    questionLine.style.font.bold = True
    questionLine.paragraph_format.line_spacing = 1

    # prepare runs line (list of choices)
    parChoices = document.add_paragraph()
    runChoices = parChoices.add_run()
    runChoices.font.size = Pt(14)
    runChoices.font.bold = False

    # for every choice and their attributes
    for choiceDict in question["choices"]:  
        create_choice_inform(runChoices, choiceDict)
        

# main code

document = Document()
styles = document.styles["Normal"]
styles.font.name = "Angsana New"

# set up page structures (border width)
sections = document.sections
for sec in sections:
    sec.top_margin = Cm(0.5)
    sec.bottom_margin = Cm(0.5)
    sec.left_margin = Cm(1)
    sec.right_margin = Cm(1)

# handle unknow heading
headingList = []

# open template(.json file)
with open('../template/template.json', 'r', encoding="utf-8") as file:
    jsonContent = json.load(file)

# extract section and sub-section data
sectionProperties = jsonContent["sectionProperties"]
templateName = jsonContent["templateName"]
lastModified = jsonContent["lastModified"]

# for every question in template file
for question in jsonContent["questions"]:
    # extract section index
    sectionNo = question["section"]
    subSectionNo = question["subSection"]
    questionType = question["questionType"]
    section = sectionProperties[sectionNo-1]
    sectionName = section["sectionName"]
    subSection = section["subSection"][subSectionNo - 1]
    
    # if found the question that is in unknow section, create new section 
    if sectionNo not in headingList:
        headingList.append(sectionNo)
        document.add_heading(f"ส่วนที่ {sectionNo}: {sectionName}", level=0)

    # add question to the form, behaviour may varies from type of question
    create_question_in_form(question)


document.save(f"../wordDocument/{templateName}.docx")


